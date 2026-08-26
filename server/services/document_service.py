from pathlib import Path

from pypdf import PdfReader
from docx import Document as DocxDocument

from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS


# =========================================================
# DIRECTORIES
# =========================================================

BASE_DIR = Path(__file__).resolve().parent.parent

UPLOAD_DIR = BASE_DIR / "data" / "uploads"

VECTORSTORE_DIR = (
    BASE_DIR / "vectorstore" / "uploads"
)


UPLOAD_DIR.mkdir(
    parents=True,
    exist_ok=True
)

VECTORSTORE_DIR.mkdir(
    parents=True,
    exist_ok=True
)


# =========================================================
# EMBEDDINGS
# =========================================================

embedding_model = HuggingFaceEmbeddings(
    model_name="sentence-transformers/all-MiniLM-L6-v2"
)


# =========================================================
# EXTRACT TEXT
# =========================================================

def extract_text(
    file_path: str,
    extension: str
):

    extension = extension.lower()

    # -----------------------------
    # PDF
    # -----------------------------

    if extension == ".pdf":

        reader = PdfReader(file_path)

        text = ""

        for page in reader.pages:

            page_text = page.extract_text()

            if page_text:
                text += page_text + "\n"

        return text

    # -----------------------------
    # DOCX
    # -----------------------------

    elif extension == ".docx":

        document = DocxDocument(
            file_path
        )

        text = "\n".join(
            paragraph.text
            for paragraph in document.paragraphs
            if paragraph.text.strip()
        )

        return text

    # -----------------------------
    # TXT
    # -----------------------------

    elif extension == ".txt":

        with open(
            file_path,
            "r",
            encoding="utf-8",
            errors="ignore"
        ) as file:

            return file.read()

    else:

        raise ValueError(
            "Unsupported file type"
        )


# =========================================================
# PROCESS DOCUMENT
# =========================================================

def process_document(
    file_path: str,
    extension: str,
    document_id: str
):

    text = extract_text(
        file_path,
        extension
    )

    if not text.strip():

        raise ValueError(
            "Could not extract text from document"
        )


    # =====================================================
    # SPLIT TEXT
    # =====================================================

    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000,
        chunk_overlap=200
    )


    chunks = text_splitter.split_documents(
        [
            Document(
                page_content=text,
                metadata={
                    "source": Path(
                        file_path
                    ).name
                }
            )
        ]
    )


    # =====================================================
    # CREATE FAISS
    # =====================================================

    vectorstore = FAISS.from_documents(
        chunks,
        embedding_model
    )


    # =====================================================
    # SAVE
    # =====================================================

    document_vectorstore_path = (
        VECTORSTORE_DIR / document_id
    )


    vectorstore.save_local(
        str(document_vectorstore_path)
    )


    return {
        "document_id": document_id,
        "chunks": len(chunks)
    }


# =========================================================
# LOAD USER DOCUMENT VECTORSTORE
# =========================================================

def load_document_vectorstore(
    document_id: str
):

    vectorstore_path = (
        VECTORSTORE_DIR / document_id
    )


    if not vectorstore_path.exists():

        raise FileNotFoundError(
            "Document vectorstore not found"
        )


    vectorstore = FAISS.load_local(
        str(vectorstore_path),
        embedding_model,
        allow_dangerous_deserialization=True
    )


    return vectorstore